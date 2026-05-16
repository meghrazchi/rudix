from __future__ import annotations

from io import BytesIO

import fitz
import pytest
from docx import Document as DocxDocument

from app.domains.documents.services.text_extraction import extract_text_sections


def _build_pdf_bytes(*, pages: list[str]) -> bytes:
    document = fitz.open()
    for text in pages:
        page = document.new_page()
        if text:
            page.insert_text((72, 72), text)
    payload = document.tobytes()
    document.close()
    return payload


def _build_docx_bytes(*, include_content: bool) -> bytes:
    doc = DocxDocument()
    if include_content:
        doc.add_paragraph("Paragraph content")
        table = doc.add_table(rows=2, cols=2)
        table.cell(0, 0).text = "H1"
        table.cell(0, 1).text = "H2"
        table.cell(1, 0).text = "R1"
        table.cell(1, 1).text = "R2"
    buffer = BytesIO()
    doc.save(buffer)
    return buffer.getvalue()


def test_extract_pdf_preserves_page_numbers() -> None:
    content = _build_pdf_bytes(pages=["Page One", "Page Two"])
    sections = extract_text_sections(file_type="pdf", content=content)

    assert [section.page_number for section in sections] == [1, 2]
    assert sections[0].text == "Page One"
    assert sections[1].text == "Page Two"
    assert sections[0].char_count == len("Page One")
    assert sections[1].char_count == len("Page Two")


def test_extract_txt_uses_utf8_fallback_for_invalid_bytes() -> None:
    content = b"hello\xffworld"
    sections = extract_text_sections(file_type="txt", content=content)

    assert len(sections) == 1
    assert sections[0].page_number == 1
    assert "hello" in sections[0].text
    assert "world" in sections[0].text
    assert sections[0].char_count == len(sections[0].text)


def test_extract_docx_includes_paragraphs_and_tables() -> None:
    content = _build_docx_bytes(include_content=True)
    sections = extract_text_sections(file_type="docx", content=content)

    assert len(sections) == 2
    assert sections[0].page_number == 1
    assert sections[0].text == "Paragraph content"
    assert sections[0].char_count == len("Paragraph content")
    assert sections[1].page_number == 2
    assert sections[1].text == "H1 | H2\nR1 | R2"
    assert sections[1].char_count == len("H1 | H2\nR1 | R2")


@pytest.mark.parametrize(
    ("file_type", "content", "expected_error"),
    [
        ("pdf", b"not-a-pdf", "malformed pdf file"),
        ("docx", b"not-a-docx", "malformed docx file"),
        ("txt", b" \n\t ", "extracted document contains no text"),
    ],
)
def test_extract_text_sections_rejects_invalid_or_empty_inputs(
    file_type: str,
    content: bytes,
    expected_error: str,
) -> None:
    with pytest.raises(ValueError, match=expected_error):
        extract_text_sections(file_type=file_type, content=content)


def test_extract_docx_rejects_empty_content() -> None:
    content = _build_docx_bytes(include_content=False)
    with pytest.raises(ValueError, match="extracted document contains no text"):
        extract_text_sections(file_type="docx", content=content)
